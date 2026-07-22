from __future__ import annotations

import csv
import os
import re
from pathlib import Path
from statistics import mean

from lxml import etree
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "revision_package"
ASSETS = PACKAGE / "response_evidence"
OUTPUT_DOCX = PACKAGE / "07_response_to_editor_and_reviewers_revised.docx"
OUTPUT_MD = PACKAGE / "08_response_to_reviewers_chinese_translation.md"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "202124"
MUTED = "5F6B7A"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "EAF3F8"
LIGHT_RED = "FDEEEE"
WHITE = "FFFFFF"

MATHML_NS = "http://www.w3.org/1998/Math/MathML"
MATH_MARKER_RE = re.compile(r"\{\{([A-Za-z0-9_]+)\}\}")
MML2OMML_PATH = Path(
    os.environ.get(
        "MML2OMML_XSL",
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
    )
)
_MML2OMML = etree.XSLT(etree.parse(str(MML2OMML_PATH)))


def mnode(tag, *children, text=None, **attrs):
    node = etree.Element(f"{{{MATHML_NS}}}{tag}")
    for key, value in attrs.items():
        node.set(key.replace("_", "-"), str(value))
    if text is not None:
        node.text = str(text)
    for child in children:
        node.append(child)
    return node


def mi(text, variant=None):
    attrs = {"mathvariant": variant} if variant else {}
    return mnode("mi", text=text, **attrs)


def mn(text):
    return mnode("mn", text=text)


def mo(text):
    return mnode("mo", text=text)


def mrow(*children):
    return mnode("mrow", *children)


def msub(base, sub):
    return mnode("msub", base, sub)


def msup(base, sup):
    return mnode("msup", base, sup)


def msubsup(base, sub, sup):
    return mnode("msubsup", base, sub, sup)


def mfrac(num, den, line=True):
    return mnode("mfrac", num, den, linethickness="1" if line else "0")


def mfenced(content, open_char="(", close_char=")"):
    return mnode("mfenced", content, open=open_char, close=close_char)


def mhat(symbol):
    return mnode("mover", symbol, mo("^"), accent="true")


def script(letter):
    return mi(letter, variant="script")


def var(letter, sub=None, sup=None):
    node = mi(letter)
    if sub is not None:
        node = msub(node, mi(sub))
    if sup is not None:
        node = msup(node, mi(sup))
    return node


def roman_sub(text):
    return mi(text, variant="normal")


def indexed(letter, sub_parts):
    sub = mrow(*sub_parts) if isinstance(sub_parts, (list, tuple)) else sub_parts
    return msub(mi(letter), sub)


def binom(top, bottom):
    return mfenced(mfrac(mn(top), bottom if hasattr(bottom, "tag") else mn(bottom), line=False))


def indiv_entry():
    return mrow(mi("indiv", variant="normal"), mo("["), mi("j"), mo(","), mi("w"), mo("]"))


def formula_joint_count():
    inner_sum = msubsup(
        mi("∑", variant="normal"),
        mrow(mi("r"), mo("="), mn("0")),
        mn("2"),
    )
    summed = mrow(inner_sum, binom("4", mi("r")))
    return mrow(
        binom("6", "3"),
        mo("×"),
        msup(mfenced(summed, "[", "]"), mn("3")),
        mo("="),
        mn("26,620"),
    )


def formula_capacity_z():
    return mrow(
        msub(mi("∑", variant="normal"), mi("w")),
        var("z", "jw"),
        mo("≤"),
        var("V", "j"),
    )


def formula_capacity_indiv():
    return mrow(
        msub(mi("∑", variant="normal"), mrow(mi("w"), mo("∈"), script("S"))),
        indiv_entry(),
        mo("≤"),
        var("V", "j"),
        mo(","),
        mo("∀"),
        mi("j"),
        mo("∈"),
        script("M"),
    )


def formula_capacity_indiv_exceed():
    return mrow(
        msub(mi("∑", variant="normal"), mi("w")),
        indiv_entry(),
        mo(">"),
        var("V", "j"),
    )


def formula_q_norm(symbol):
    return mrow(
        mhat(mi(symbol)),
        mo("="),
        mfrac(
            mrow(mi(symbol), mo("−"), indexed(symbol, roman_sub("min"))),
            mrow(
                indexed(symbol, roman_sub("max")),
                mo("−"),
                indexed(symbol, roman_sub("min")),
            ),
        ),
    )


def formula_q_score():
    return mrow(
        mi("Q"),
        mo("="),
        mi("λ"),
        mhat(mi("C")),
        mo("+"),
        mfenced(mrow(mn("1"), mo("−"), mi("λ"))),
        mhat(mi("D")),
        mo(","),
        mi("λ"),
        mo("="),
        mn("0.5"),
    )


def formula_varpi_rule():
    return mrow(
        var("ϖ", "j"),
        mo("="),
        mo("⌈"),
        mi("ρ"),
        var("V", "j"),
        mo("⌉"),
        mo(","),
        mi("ρ"),
        mo("="),
        mn("0.5"),
    )


def formula_hybrid_score():
    return mrow(
        var("H", "j"),
        mfenced(var("s", "w")),
        mo("="),
        mi("α"),
        indexed("w", [roman_sub("cost"), mo(","), mi("j")]),
        mfenced(var("s", "w")),
        mo("+"),
        mi("β"),
        indexed("w", [roman_sub("req"), mo(","), mi("j")]),
        mfenced(var("s", "w")),
    )


def formula_pm_eta():
    return mrow(
        var("p", "m"),
        mo("="),
        mfrac(
            mn("1"),
            mrow(mi("k"), mo("|"), script("S"), mo("|")),
        ),
        mo(","),
        var("η", "m"),
        mo("="),
        mn("20"),
    )


FORMULA_BUILDERS = {
    "x_ijw": lambda: var("x", "ijw"),
    "x_jw": lambda: var("x", "jw"),
    "x_jk": lambda: var("x", "jk"),
    "y_jk": lambda: var("y", "jk"),
    "z_jw": lambda: var("z", "jw"),
    "a_ijw": lambda: var("a", "ijw"),
    "w_j": lambda: var("w", "j"),
    "m_j": lambda: var("m", "j"),
    "b_k": lambda: var("b", "k"),
    "u_i": lambda: var("u", "i"),
    "s_w": lambda: var("s", "w"),
    "s_1": lambda: var("s", "1"),
    "s_2": lambda: var("s", "2"),
    "s_3": lambda: var("s", "3"),
    "s_4": lambda: var("s", "4"),
    "s_5": lambda: var("s", "5"),
    "s_6": lambda: var("s", "6"),
    "V_j": lambda: var("V", "j"),
    "D_i": lambda: var("D", "i"),
    "Nsym": lambda: mi("N"),
    "Gsym": lambda: mi("G"),
    "Qsym": lambda: mi("Q"),
    "varpi_j": lambda: var("ϖ", "j"),
    "indiv_jw": indiv_entry,
    "joint_count": formula_joint_count,
    "capacity_z": formula_capacity_z,
    "capacity_indiv": formula_capacity_indiv,
    "capacity_indiv_exceed": formula_capacity_indiv_exceed,
    "q_norm_C": lambda: formula_q_norm("C"),
    "q_norm_D": lambda: formula_q_norm("D"),
    "q_score": formula_q_score,
    "varpi_rule": formula_varpi_rule,
    "hybrid_score": formula_hybrid_score,
    "N50": lambda: mrow(mi("N"), mo("="), mn("50")),
    "G200": lambda: mrow(mi("G"), mo("="), mn("200")),
    "pc_eta": lambda: mrow(var("p", "c"), mo("="), mn("0.9"), mo(","), var("η", "c"), mo("="), mn("15")),
    "pm_eta": formula_pm_eta,
    "alpha_beta": lambda: mrow(mi("α"), mo("="), mi("β"), mo("="), mn("0.5")),
    "Vj4": lambda: mrow(var("V", "j"), mo("="), mn("4")),
    "rho05": lambda: mrow(mi("ρ"), mo("="), mn("0.5")),
    "varpi2": lambda: mrow(var("ϖ", "j"), mo("="), mn("2")),
    "varpi4": lambda: mrow(var("ϖ", "j"), mo("="), mn("4")),
    "alpha_beta_range": lambda: mrow(mi("α"), mo(","), mi("β"), mo("∈"), mfenced(mrow(mn("0"), mo(","), mn("1")), "[", "]"), mo(","), mi("α"), mo("+"), mi("β"), mo("="), mn("1")),
    "varpi_bounds": lambda: mrow(mn("0"), mo("<"), var("ϖ", "j"), mo("≤"), var("V", "j")),
    "Q03282": lambda: mrow(mi("Q"), mo("="), mn("0.3282")),
    "Q06125": lambda: mrow(mi("Q"), mo("="), mn("0.6125")),
    "HV09470": lambda: mrow(mi("HV", variant="normal"), mo("="), mn("0.9470")),
    "IGD00016": lambda: mrow(mi("IGD", variant="normal"), mo("="), mn("0.0016")),
    "HV10116": lambda: mrow(mi("HV", variant="normal"), mo("="), mn("1.0116")),
    "IGD0129": lambda: mrow(mi("IGD", variant="normal"), mo("="), mn("0.0129")),
    "Q02678": lambda: mrow(mi("Q"), mo("="), mn("0.2678")),
    "Q05517": lambda: mrow(mi("Q"), mo("="), mn("0.5517")),
    "learning_rate": lambda: mrow(mn("7"), mo("×"), msup(mn("10"), mrow(mo("−"), mn("4")))),
    "preference_weights": lambda: mfenced(mrow(mn("0.1"), mo(","), mn("0.3"), mo(","), mn("0.5"), mo(","), mn("0.7"), mo(","), mn("0.9")), "{", "}"),
    "service_range": lambda: mrow(var("s", "1"), mo("–"), var("s", "6")),
    "indiv_matrix": lambda: mrow(
        mi("indiv", variant="normal"),
        mo("∈"),
        msup(
            mfenced(mrow(mn("0"), mo(","), mn("1")), "{", "}"),
            mrow(mi("k"), mo("×"), mo("|"), script("S"), mo("|")),
        ),
    ),
    "user_set": lambda: mrow(script("U"), mo("="), mfenced(var("u", "i"), "{", "}")),
    "tilde_D_i": lambda: mnode("mover", var("D", "i"), mo("~"), accent="true"),
    "ref_point": lambda: mfenced(mrow(mn("1.1"), mo(","), mn("1.1"))),
    "w_cost_j": lambda: indexed("w", [roman_sub("cost"), mo(","), mi("j")]),
    "w_req_j": lambda: indexed("w", [roman_sub("req"), mo(","), mi("j")]),
}

FORMULA_LATEX = {
    "x_ijw": r"x_{ijw}",
    "x_jw": r"x_{jw}",
    "x_jk": r"x_{jk}",
    "y_jk": r"y_{jk}",
    "z_jw": r"z_{jw}",
    "a_ijw": r"a_{ijw}",
    "w_j": r"w_j",
    "m_j": r"m_j",
    "b_k": r"b_k",
    "u_i": r"u_i",
    "s_w": r"s_w",
    "s_1": r"s_1",
    "s_2": r"s_2",
    "s_3": r"s_3",
    "s_4": r"s_4",
    "s_5": r"s_5",
    "s_6": r"s_6",
    "V_j": r"V_j",
    "D_i": r"D_i",
    "Nsym": r"N",
    "Gsym": r"G",
    "Qsym": r"Q",
    "varpi_j": r"\varpi_j",
    "indiv_jw": r"\mathrm{indiv}[j,w]",
    "joint_count": r"\binom{6}{3}\times\left[\sum_{r=0}^{2}\binom{4}{r}\right]^3=26{,}620",
    "capacity_z": r"\sum_w z_{jw}\leq V_j",
    "capacity_indiv": r"\sum_{w\in\mathcal{S}}\mathrm{indiv}[j,w]\leq V_j,\ \forall j\in\mathcal{M}",
    "capacity_indiv_exceed": r"\sum_w\mathrm{indiv}[j,w]>V_j",
    "q_norm_C": r"\widehat{C}=\frac{C-C_{\min}}{C_{\max}-C_{\min}}",
    "q_norm_D": r"\widehat{D}=\frac{D-D_{\min}}{D_{\max}-D_{\min}}",
    "q_score": r"Q=\lambda\widehat{C}+(1-\lambda)\widehat{D},\ \lambda=0.5",
    "varpi_rule": r"\varpi_j=\lceil\rho V_j\rceil,\ \rho=0.5",
    "hybrid_score": r"H_j(s_w)=\alpha w_{\mathrm{cost},j}(s_w)+\beta w_{\mathrm{req},j}(s_w)",
    "N50": r"N=50",
    "G200": r"G=200",
    "pc_eta": r"p_c=0.9,\ \eta_c=15",
    "pm_eta": r"p_m=\frac{1}{k|\mathcal{S}|},\ \eta_m=20",
    "alpha_beta": r"\alpha=\beta=0.5",
    "Vj4": r"V_j=4",
    "rho05": r"\rho=0.5",
    "varpi2": r"\varpi_j=2",
    "varpi4": r"\varpi_j=4",
    "alpha_beta_range": r"\alpha,\beta\in[0,1],\ \alpha+\beta=1",
    "varpi_bounds": r"0<\varpi_j\leq V_j",
    "Q03282": r"Q=0.3282",
    "Q06125": r"Q=0.6125",
    "HV09470": r"\mathrm{HV}=0.9470",
    "IGD00016": r"\mathrm{IGD}=0.0016",
    "HV10116": r"\mathrm{HV}=1.0116",
    "IGD0129": r"\mathrm{IGD}=0.0129",
    "Q02678": r"Q=0.2678",
    "Q05517": r"Q=0.5517",
    "learning_rate": r"7\times10^{-4}",
    "preference_weights": r"\{0.1,0.3,0.5,0.7,0.9\}",
    "service_range": r"s_1\text{--}s_6",
    "indiv_matrix": r"\mathrm{indiv}\in\{0,1\}^{k\times|\mathcal{S}|}",
    "user_set": r"\mathcal{U}=\{u_i\}",
    "tilde_D_i": r"\widetilde{D_i}",
    "ref_point": r"(1.1,1.1)",
    "w_cost_j": r"w_{\mathrm{cost},j}",
    "w_req_j": r"w_{\mathrm{req},j}",
}


def formula_omml(key):
    if key not in FORMULA_BUILDERS:
        raise KeyError(f"Unknown formula marker: {key}")
    math = etree.Element(f"{{{MATHML_NS}}}math", nsmap={None: MATHML_NS})
    math.append(FORMULA_BUILDERS[key]())
    transformed = _MML2OMML(math)
    return parse_xml(etree.tostring(transformed, encoding="unicode"))


def add_rich_text(paragraph, text, size=11, bold=False, italic=False, color=TEXT):
    cursor = 0
    for match in MATH_MARKER_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, bold=bold, italic=italic, color=color)
        paragraph._p.append(formula_omml(match.group(1)))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def markdown_math(text):
    return MATH_MARKER_RE.sub(lambda match: f"${FORMULA_LATEX[match.group(1)]}$", text)


def set_run_font(run, size=11, bold=False, italic=False, color=TEXT, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D0D5DD", size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    node.set(qn("w:val"), "1" if value else "0")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color, before, after in (
        ("Title", 18, DARK_BLUE, 0, 10),
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("MOS\u00b2 | Response to the Editor and Reviewers")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Response to the Editor and Reviewers")
    set_run_font(r, size=18, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Original manuscript ID: IoT-65990-2026")
    set_run_font(r, size=10.5, bold=True, color=TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("MOS\u00b2: A Two-Stage Multi-Objective Framework for Server Deployment and Service Provisioning in Mobile Edge Computing")
    set_run_font(r, size=11, italic=True, color=TEXT)

    table = doc.add_table(rows=1, cols=1)
    keep_row_together(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.55)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, color="B8D4E3", size=6)
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Revision format. ")
    set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)
    r = p.add_run("All changes are highlighted in blue in the marked manuscript; a clean manuscript is provided separately.")
    set_run_font(r, size=9.5, color=TEXT)


def add_body_paragraph(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        add_rich_text(p, text[len(bold_prefix):], italic=italic)
    else:
        add_rich_text(p, text, italic=italic)
    return p


def add_comment_box(doc, text, label="Reviewer Comment"):
    table = doc.add_table(rows=1, cols=1)
    keep_row_together(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.55)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_border(cell, color="CFD4DC", size=6)
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_rich_text(p, text, size=10, italic=True, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_response_label(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Response")
    set_run_font(r, size=10.5, bold=True, color=BLUE)
    keep_with_next(p)


def add_placeholder(doc, text):
    table = doc.add_table(rows=1, cols=1)
    keep_row_together(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.55)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, color="7FAAC2", size=8)
    set_cell_margins(cell, top=90, bottom=90, start=130, end=130)
    p = cell.paragraphs[0]
    r = p.add_run("MANUSCRIPT EXCERPT PLACEHOLDER: ")
    set_run_font(r, size=9.2, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=9.2, italic=True, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_revised_text_box(doc, excerpt):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    keep_row_together(header)
    cell = header.cells[0]
    cell.width = Inches(6.55)
    set_cell_shading(cell, DARK_BLUE)
    set_cell_border(cell, color=DARK_BLUE, size=7)
    set_cell_margins(cell, top=75, bottom=75, start=125, end=125)
    p = cell.paragraphs[0]
    r = p.add_run("Revised Manuscript Text")
    set_run_font(r, size=9.6, bold=True, color=WHITE)
    keep_with_next(p)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Location: {excerpt['location']}")
    set_run_font(r, size=8.8, italic=True, color=WHITE)
    keep_with_next(p)

    for block in excerpt["blocks"]:
        row = table.add_row()
        keep_row_together(row)
        cell = row.cells[0]
        cell.width = Inches(6.55)
        set_cell_shading(cell, "F5FAFD")
        set_cell_border(cell, color="8AB7CF", size=5)
        set_cell_margins(cell, top=75, bottom=75, start=125, end=125)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if block["type"] == "equation":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p._p.append(formula_omml(block["key"]))
        else:
            add_rich_text(
                p,
                block["en"],
                size=9.35,
                italic=block["type"] == "caption",
                color=TEXT,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, path, caption, width=5.8):
    if not Path(path).exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    keep_with_next(p)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(7)
    r = cap.add_run(caption)
    set_run_font(r, size=9, italic=True, color=MUTED)
    return cap


def add_data_table(doc, headers, rows, widths=None, bold_cells=None, font_size=9.0, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(caption)
        set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)
        keep_with_next(p)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    bold_cells = set(bold_cells or [])
    if widths is None:
        widths = [6.45 / len(headers)] * len(headers)
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.width = Inches(widths[j])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, DARK_BLUE)
        set_cell_border(cell, color="AEB7C2", size=5)
        set_cell_margins(cell, top=70, bottom=70, start=80, end=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rich_text(p, str(header), size=font_size, bold=True, color=WHITE)
    repeat_table_header(table.rows[0])
    for i, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cell = cells[j]
            cell.width = Inches(widths[j])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, WHITE if i % 2 else "F8FAFC")
            set_cell_border(cell, color="D0D5DD", size=4)
            set_cell_margins(cell, top=60, bottom=60, start=75, end=75)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, str(value), size=font_size, bold=(i - 1, j) in bold_cells)
    # These response tables are intentionally short. Keeping their rows together
    # avoids orphaned headers or a single row stranded at a page boundary.
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                keep_with_next(paragraph)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_equation(doc, key):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p._p.append(formula_omml(key))


def read_csv(name):
    with (ASSETS / name).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def joint_rows():
    data = read_csv("joint_gap_summary.csv")
    rows = []
    for d in data:
        seed = d["Config"].split("seed")[-1]
        rows.append([
            seed,
            f'{float(d["HVGapPct"]):.2f}%',
            f'{float(d["IGDGap"]):.4f}',
            f'{float(d["BestQGap"]):.4f}',
            f'{float(d["RuntimeRatioJointOverMOS2"]):.2f}x',
        ])
    rows.append([
        "Mean",
        f'{mean(float(d["HVGapPct"]) for d in data):.2f}%',
        f'{mean(float(d["IGDGap"]) for d in data):.4f}',
        f'{mean(float(d["BestQGap"]) for d in data):.4f}',
        f'{mean(float(d["RuntimeRatioJointOverMOS2"]) for d in data):.2f}x',
    ])
    return rows


def real_region_stage2_rows():
    data = read_csv("real_region_stage2_aggregate.csv")
    return [[
        d["Method"],
        f'{float(d["HVMean"]):.4f} \u00b1 {float(d["HVStd"]):.4f}',
        f'{float(d["IGDMean"]):.4f} \u00b1 {float(d["IGDStd"]):.4f}',
        f'{float(d["BestQMean"]):.4f} \u00b1 {float(d["BestQStd"]):.4f}',
    ] for d in data]


def real_region_bestq_rows():
    data = read_csv("real_region_bestq_aggregate.csv")
    return [[d["Method"], f'{float(d["BestQMean"]):.4f}', f'{float(d["BestQStd"]):.4f}'] for d in data]


def stage2_dqn_rows():
    data = read_csv("stage2_bestq_with_dqn.csv")
    index = {(d["Config"], d["Method"]): float(d["BestQ"]) for d in data}
    configs = [
        ("Fixed 10 servers", "10_100"),
        ("Fixed 10 servers", "10_130"),
        ("Fixed 10 servers", "10_150"),
        ("Fixed 10 servers", "10_180"),
        ("Fixed 130 users", "5_130"),
        ("Fixed 130 users", "10_130"),
        ("Fixed 130 users", "15_130"),
        ("Fixed 130 users", "20_130"),
    ]
    rows = []
    for series, config in configs:
        psp = index[(config, "PSP")]
        dqn = index[(config, "DQN")]
        reduction = 100.0 * (dqn - psp) / dqn
        rows.append([series, config.replace("_", "/"), f"{psp:.4f}", f"{dqn:.4f}", f"{reduction:.1f}%"])
    return rows


ITEMS = [
    {
        "section": "Associate Editor",
        "number": "AE",
        "title": "Overall assessment and reference selection",
        "comment_en": (
            "This is an extended version of a conference paper. I can see there are quite some new contents for algorithm, experiments, etc., so a submission to a journal is reasonable to me. "
            "We got the comments from two reviewers, both offered detailed comments and overall recommendations, that are not positive enough for an acceptance. "
            "I recommend a reject, but meanwhile suggest to offer the authors a chance to revise significantly and re-submit; however the authors can also consider submitting this work to other more suitable journals. "
            "Note that the authors shall feel free to evaluate each reference and only cite those with true and big enough relevance to this study."
        ),
        "comment_zh": (
            "这是会议论文的扩展版本。我看到在算法、实验等方面加入了相当多的新内容，因此向期刊投稿是合理的。两位审稿人都给出了详细意见，但总体建议尚不足以支持接收。副编辑建议拒稿，同时给予作者进行大幅修改后重新投稿的机会；作者也可以考虑其他更合适的期刊。作者可自行评估每篇建议文献，只引用与本研究真正且充分相关的工作。"
        ),
        "response_en": [
            "Thank you for recognizing the journal-level extension and for allowing a substantially revised resubmission. We have addressed every technical and presentation issue raised by the two reviewers. The marked manuscript highlights all changes in blue, and a clean version is provided separately. The revision includes a redesigned system illustration and algorithm flow, a unified notation system, a technical rationale for the two-stage formulation, an explicit feasibility-repair procedure, additional reproducibility details, a CLS initialization-sensitivity study, a DQN baseline, quantitative Pareto metrics, and focused supporting experiments.",
            "We also followed the guidance on references. Each of the six suggested works was evaluated individually. Four directly relevant studies on service deployment/offloading, local search, long-term MEC resource allocation, and online learning were incorporated into the Related Work section. The two FSO-enabled SAGIN task-offloading studies were not cited because their network architecture, link model, and decision variables differ substantially from the terrestrial capacity-constrained service-provisioning problem considered here."
        ],
        "response_zh": [
            "感谢副编辑认可本稿作为期刊扩展版本的合理性，并给予大幅修改后重新投稿的机会。我们已经逐条处理两位审稿人提出的技术与表达问题。标记稿以蓝色显示全部修改，并另附无标记正式稿。主要修改包括重构系统示意图和算法流程图、统一符号体系、补充两阶段分解的技术依据、明确可行性修复流程、完善复现参数、增加 CLS 初始化敏感性实验、DQN 学习型基线、Pareto 定量指标以及有针对性的补充实验。",
            "我们也遵循了参考文献选择建议，对六篇推荐文献逐篇评估。与服务部署/卸载、局部搜索、长期 MEC 资源分配和在线学习直接相关的四篇文献已纳入 Related Work；另外两篇聚焦 FSO 支持的 SAGIN 任务卸载，其网络架构、链路模型和决策变量与本文的地面 MEC 容量受限服务配置问题差异较大，因此未机械引用。"
        ],
    },
    {
        "section": "Reviewer 1",
        "number": "1",
        "title": "Clarity of Fig. 1 and the dashed regions",
        "comment_en": (
            "Figure 1 in the paper is somewhat cluttered and difficult to follow. Specifically, the regions enclosed by the dashed lines are not clearly explained. "
            "The authors should refine the visual layout of this figure and provide a more explicit description in both the text and the caption regarding what these dashed areas represent."
        ),
        "comment_zh": "论文中的图 1 有些拥挤，不易理解。尤其是虚线圈出的区域没有得到清楚解释。作者应改进图的视觉布局，并在正文和图注中更明确地说明这些虚线区域代表什么。",
        "response_en": [
            "Thank you for identifying the ambiguity in Fig. 1. We revised the illustration and its accompanying explanation so that every visual element has a unique meaning. The solid black contours partition the topology into interconnected service regions; the black dashed circles denote individual base-station coverage areas; the red dotted circles identify base stations equipped with edge servers; blue bidirectional arrows represent cloud-edge communication; and red dashed arrows indicate inter-region request forwarding. User colors correspond to requested service types, and the service blocks beside an edge server identify its instantiated services. The question-mark slot in Area 3 deliberately represents the remaining provisioning decision under limited capacity.",
            "The caption is now a concise one-sentence description, while the main text explains the dashed regions, routing arrows, service colors, instantiated services, and the open service slot in detail."
        ],
        "response_zh": [
            "感谢审稿人指出图 1 的歧义。我们重新设计了图示及其配套说明，使每一种视觉元素只表达一种含义：黑色实线轮廓划分相互连接的服务区域；黑色虚线圆表示单个基站覆盖范围；红色点线圆标识配置了边缘服务器的基站；蓝色双向箭头表示云边通信；红色虚线箭头表示跨区域请求转发。用户颜色与其请求的服务类型对应，边缘服务器旁的色块表示已实例化服务，Area 3 的问号槽位用于表示容量受限条件下尚待确定的服务配置。",
            "图注已压缩为一句概括性说明，虚线区域、转发箭头、服务颜色、已配置服务和开放槽位的具体含义均在正文中详细解释。"
        ],
        "placeholder": "Insert the marked-manuscript excerpt containing revised Fig. 1 and its explanatory paragraphs in Section I-A (marked manuscript, pp. 2-3).",
        "evidence": "fig1",
    },
    {
        "section": "Reviewer 1",
        "number": "2",
        "title": "Conflicting decision-variable notation",
        "comment_en": (
            "There are some conflicts and visual similarities in the defined notations, particularly concerning the decision variables. "
            "For example, {{x_ijw}} and {{x_jw}} are used to denote different concepts, which may easily confuse the readers. "
            "It is strongly recommended to use completely distinct variable letters to differentiate the decision variables clearly."
        ),
        "comment_zh": "已定义符号中存在冲突和视觉相似，尤其是决策变量。例如 {{x_ijw}} 与 {{x_jw}} 表示不同概念，容易使读者混淆。强烈建议使用完全不同的字母清楚区分这些决策变量。",
        "response_en": [
            "We agree that the original notation could cause confusion. The decision variables have been renamed consistently throughout the system model, Table I, constraints, objective functions, and algorithm descriptions: {{y_jk}} denotes server deployment, {{z_jw}} denotes service provisioning, and {{a_ijw}} denotes user-service-server association. The implementation entry {{indiv_jw}} is explicitly defined as the encoded counterpart of {{z_jw}}. This revision removes the former visual conflict among {{x_jk}}, {{x_jw}}, and {{x_ijw}} and also eliminates the inconsistent use of {{w_j}} for server {{m_j}}."
        ],
        "response_zh": [
            "我们同意原符号体系容易造成混淆。系统模型、Table I、约束、目标函数和算法描述中的决策变量已统一重命名：{{y_jk}} 表示服务器部署，{{z_jw}} 表示服务配置，{{a_ijw}} 表示用户-服务-服务器关联。实现中的矩阵项 {{indiv_jw}} 也明确说明为 {{z_jw}} 的编码对应项。这样消除了 {{x_jk}}、{{x_jw}} 和 {{x_ijw}} 之间的视觉冲突，同时修正了将服务器 {{m_j}} 误写为 {{w_j}} 的不一致。"
        ],
        "placeholder": "Insert the marked-manuscript excerpts from Section III-A and Table I showing y_jk, z_jw, and a_ijw (marked manuscript, pp. 4-5).",
        "evidence": "notation",
    },
    {
        "section": "Reviewer 1",
        "number": "3",
        "title": "Rationale and potential optimality loss of the two-stage decomposition",
        "comment_en": (
            "The authors transform the highly intertwined server deployment and service provisioning problem into two relatively independent stages. "
            "While this decoupling effectively reduces computational complexity, it raises the question of whether this transformation leads to a loss of global optimality. "
            "The authors should add a discussion justifying this two-stage decomposition and, if possible, comment on or analyze the potential performance gap compared to a joint optimization approach."
        ),
        "comment_zh": "作者将高度耦合的服务器部署与服务配置问题转化为两个相对独立的阶段。虽然这种解耦降低了计算复杂度，但也引出了是否损失全局最优性的问题。作者应说明两阶段分解的合理性，并在可能的情况下分析其相对联合优化方法的性能差距。",
        "response_en": [
            "Thank you for raising this important point. The revised manuscript explains that the decomposition follows the operational hierarchy of MEC planning. Server deployment is a relatively long-term infrastructure decision governed by budget, coverage, and geographic demand, whereas service instances are provisioned after physical server locations are known. Stage I passes the selected locations, user assignments, and deployment-related cost to Stage II, thereby preserving the principal cost-latency dependence while substantially reducing the combinatorial decision space.",
            "We further conducted a response-only exact small-scale joint-optimization experiment. The instance contains 6 candidate stations, 3 deployed servers, 30 users, 4 service types, and a per-server service capacity of 2. All {{joint_count}} feasible joint server-service decisions were enumerated to construct an exact Pareto reference. Across seeds 42, 43, and 44, MOS\u00b2-PSP attained the same best normalized weighted quality as the exact joint reference in every run. The mean HV gap was 3.87%, the mean IGD was 0.0382, and exact enumeration required 5.18 times the MOS\u00b2 runtime on average even at this reduced scale.",
            "The near-identical balanced cost-delay quality and the steep growth of exhaustive search show that decomposition is not merely a computational convenience, but is necessary for tractable optimization at the larger MEC scales studied in the manuscript. MOS\u00b2 preserves the decisive deployment-to-provisioning dependencies while matching the exact solver's best balanced solution in every tested seed, using only about one-fifth of its runtime and avoiding enumeration and storage of the full joint decision space."
        ],
        "response_zh": [
            "感谢审稿人提出这一关键问题。修改稿说明，两阶段分解遵循 MEC 规划的实际决策层级：服务器部署是受预算、覆盖和地理需求约束的长期基础设施决策，服务实例则在物理服务器位置确定后进行配置。Stage I 将已选位置、用户分配和部署相关成本传递给 Stage II，因此在显著缩小组合决策空间的同时保留主要的成本-时延联系。",
            "我们还进行了仅用于回信佐证的小规模精确联合优化实验。实例包含 6 个候选基站、3 台部署服务器、30 个用户、4 类服务，每台服务器服务容量为 2。通过枚举全部 {{joint_count}} 个可行联合决策，构造精确 Pareto 参考。在种子 42、43 和 44 下，MOS\u00b2-PSP 每次都取得与精确联合参考相同的最佳归一化加权质量；平均 HV 差距为 3.87%，平均 IGD 为 0.0382，而精确枚举即使在该小规模下平均也需要 MOS\u00b2 的 5.18 倍运行时间。",
            "近乎一致的成本-时延平衡质量与穷举搜索负担的陡增表明，两阶段分解并非单纯为了计算方便，而是使本文较大规模 MEC 配置可求解的必要设计。MOS\u00b2 在保留部署与服务配置关键依赖关系的同时，在全部测试种子下均取得与精确求解器相同的最佳平衡解，运行时间仅约为后者的五分之一，并避免枚举和存储完整联合决策空间。"
        ],
        "placeholder": "Insert the marked-manuscript paragraph that explains the operational rationale of the two-stage decomposition in Section III-D (marked manuscript, p. 6). The exact joint experiment below is supporting evidence for this response and is not inserted into the manuscript.",
        "evidence": "joint",
    },
    {
        "section": "Reviewer 1",
        "number": "4",
        "title": "Constraint checking and handling in PSP",
        "comment_en": (
            "The PSP strategy was proposed for the service provisioning stage, the logic surrounding the constraint check step is somewhat vague and unclear. "
            "The authors need to adjust this part of the diagram and elaborate in the corresponding text on exactly how the constraints are verified and handled during the algorithm's execution (e.g., whether penalty functions or repair mechanisms are used)."
        ),
        "comment_zh": "PSP 用于服务配置阶段，但约束检查步骤周围的逻辑较为模糊。作者需要调整图中这一部分，并在正文中说明算法执行期间如何验证和处理约束，例如使用惩罚函数还是修复机制。",
        "response_en": [
            "Fig. 2, Algorithm 3, and the associated text now describe the complete feasibility procedure. After crossover and mutation, each decision entry is rounded to a binary value. For every server {{m_j}}, PSP checks whether {{capacity_z}}. If capacity is exceeded, selected service entries are randomly deactivated until feasibility is restored. Objective evaluation and non-dominated sorting are performed only after repair. PSP therefore enforces capacity through an explicit repair mechanism rather than a penalty function. The redesigned Evolutionary Optimization panel presents the same sequence and uses proper mathematical subscripts for the population and capacity variables."
        ],
        "response_zh": [
            "Fig. 2、Algorithm 3 及相应正文现已完整描述可行性处理流程。交叉和变异后，先将每个决策项二值化；随后对每台服务器 {{m_j}} 检查 {{capacity_z}}。若容量超限，则随机关闭已选服务项，直到恢复可行。只有修复后的个体才进入目标评价和非支配排序。因此 PSP 使用的是显式修复机制，而不是惩罚函数。重绘后的 Evolutionary Optimization 面板按同一顺序展示流程，并使用规范的种群和容量数学下标。"
        ],
        "placeholder": "Insert the marked-manuscript excerpts containing Fig. 2, the PSP feasibility paragraph, and the revised repair lines in Algorithm 3 (marked manuscript, pp. 6-8).",
        "evidence": "fig2",
    },
    {
        "section": "Reviewer 1",
        "number": "5",
        "title": "Readability of axes, labels, and ticks",
        "comment_en": "The font sizes for the X-axis and Y-axis labels/ticks in all experimental figures are too small. Please enlarge them to ensure they are easily readable.",
        "comment_zh": "所有实验图中 X 轴和 Y 轴标签及刻度字号过小。请放大这些文字以确保易于阅读。",
        "response_en": [
            "We re-exported the experimental figures with consistent multi-panel dimensions, larger axis labels and tick labels, improved panel spacing, and concise captions. The Stage-I scale and convergence results, Stage-II scalar comparisons, hybrid-initialization illustration, Pareto fronts, and CLS sensitivity figure were all visually checked at the final IEEE column widths. Redundant shared legends and auxiliary annotations were removed where they reduced usable plotting area."
        ],
        "response_zh": [
            "我们重新导出了全部实验图，统一多子图尺寸，增大坐标轴标题和刻度文字，改善子图间距，并将图注压缩为简洁说明。Stage I 规模与收敛结果、Stage II 标量比较、混合初始化示意图、Pareto 前沿和 CLS 敏感性图均按最终 IEEE 栏宽进行了视觉检查；对于挤占绘图区的重复图例和辅助标注进行了精简。"
        ],
        "placeholder": "Insert representative marked-manuscript screenshots of the revised experimental figures in Section V (marked manuscript, pp. 10-13).",
        "evidence": "readability",
    },
    {
        "section": "Reviewer 1",
        "number": "6",
        "title": "Typographical and language errors",
        "comment_en": (
            "There are several typos in this paper. A thorough proofreading is required. Some specific examples include: page 5: s2 , s3 and s5. -> s2 , s3, and s5; page 12: U the set of -> U is the set of."
        ),
        "comment_zh": "论文中存在若干拼写或排版错误，需要全面校对。具体例子包括：第 5 页的 s2 , s3 and s5. 应改为 s2, s3, and s5；第 12 页的 U the set of 应改为 U is the set of。",
        "response_en": [
            "The manuscript has been proofread throughout. The two examples identified by the reviewer were corrected. We also corrected ordinal suffixes in the author affiliations, subject-verb agreement in the Introduction, the phrase 'hierarchy network architecture,' repeated introductory wording, inconsistent server symbols, punctuation in service lists, and grammatical issues in the system-model and algorithm descriptions. The marked and clean manuscripts were both compiled and visually inspected after proofreading."
        ],
        "response_zh": [
            "我们已对全文进行校对并修正审稿人指出的两处问题。此外，还修正了作者序号后缀、引言中的主谓一致、hierarchy network architecture 的词性错误、重复的引导语、服务器符号不一致、服务列表标点以及系统模型和算法描述中的语法问题。校对后，标记稿和无标记稿均重新编译并进行了视觉检查。"
        ],
        "placeholder": "Insert selected marked-manuscript excerpts showing the corrected service-list punctuation, the completed definition of U, and representative proofreading changes.",
        "evidence": "proofreading",
    },
    {
        "section": "Reviewer 2",
        "number": "1",
        "title": "Initialization sensitivity of CLS",
        "comment_en": (
            "The proposed CLS algorithm is essentially a local-search-based heuristic for the K-median problem. However, the manuscript does not clearly explain the initialization sensitivity of the algorithm. "
            "Since the initial deployment set S in Algorithm 1 is randomly generated, different initializations may lead to significantly different local optima."
        ),
        "comment_zh": "CLS 本质上是 K-median 问题的局部搜索启发式方法，但稿件没有清楚解释其初始化敏感性。由于 Algorithm 1 的初始部署集合 S 随机生成，不同初始化可能导致显著不同的局部最优解。",
        "response_en": [
            "We added a 50-run initialization-sensitivity study comparing Random, Density, Distance-Sum, marginal Greedy, and Diverse initializations. For each server/user configuration, the reported gap is the percentage difference between the final CLS cost and the best final cost observed under the same configuration.",
            "With 130 users and 5, 10, 15, or 20 deployed servers, Random reached the same best final cost as the specialized initialization strategies in three of the four configurations and showed only a small 2.10% mean gap in the remaining case. More importantly, in the complementary 10-server/150-user stress case, Random obtained a 1.27% mean gap, whereas marginal Greedy deteriorated to 15.88%. Thus, handcrafted initialization provides no consistent optimization advantage, and a deterministic greedy preference can even steer CLS toward a markedly poorer local optimum. We therefore retain Random initialization because it matches the dedicated strategies in most tested settings while remaining more robust to initialization bias."
        ],
        "response_zh": [
            "我们增加了 50 次重复的初始化敏感性实验，比较 Random、Density、Distance-Sum、marginal Greedy 和 Diverse 五种初始化。对每个服务器/用户配置，gap 定义为最终 CLS 成本相对同一配置下观测到的最佳最终成本的百分比差。",
            "在固定 130 个用户、服务器数为 5、10、15 和 20 时，Random 在四组中的三组都与专门设计的初始化策略达到相同最佳最终成本，在其余一组中也仅有 2.10% 的较小平均 gap。更重要的是，在补充的 10 台服务器/150 个用户压力情形中，Random 的平均 gap 为 1.27%，而 marginal Greedy 显著恶化至 15.88%。因此，手工设计的初始化并未带来稳定一致的优化优势，确定性的贪心偏好在特定情况下反而可能将 CLS 引向明显较差的局部最优。本文据此保留 Random 初始化：它在多数测试设置下与专门策略效果一致，同时对初始化偏置更稳健。"
        ],
        "placeholder": "Insert the marked-manuscript excerpt containing the CLS initialization-sensitivity paragraph and Fig. 9 in Section V (marked manuscript, p. 12).",
        "evidence": "cls",
    },
    {
        "section": "Reviewer 2",
        "number": "2",
        "title": "Reliability-related QoS metrics",
        "comment_en": (
            "In Eq. (6), the QoS constraint only constrains the end-to-end latency upper bound {{D_i}}. However, packet loss, reliability, and service interruption probability are not considered. "
            "Since MEC systems for latency-sensitive applications usually require reliability guarantees, the manuscript is suggested to discuss the impact of ignoring reliability-related QoS metrics."
        ),
        "comment_zh": "公式 (6) 的 QoS 约束只限制端到端时延上界 {{D_i}}，没有考虑丢包率、可靠性和服务中断概率。由于时延敏感型 MEC 应用通常需要可靠性保障，建议讨论忽略可靠性相关 QoS 指标的影响。",
        "response_en": [
            "Equation (6) retains end-to-end latency as the primary QoS requirement because the present optimization focuses on the cost-latency trade-off. The accompanying model discussion now clarifies that, for reliability-critical applications, latency compliance can be complemented by packet-loss, link-availability, and service-interruption constraints so that low-latency solutions are not selected when continuity is insufficient. The Conclusion also identifies reliability-aware QoS constraints as a direct extension of the framework."
        ],
        "response_zh": [
            "公式 (6) 保留端到端时延作为主要 QoS 要求，因为本文优化聚焦成本-时延权衡。模型部分现已说明，对于可靠性关键型应用，可在时延约束之外补充丢包率、链路可用性和服务中断约束，以避免选择连续性不足的低时延方案。结论中也将可靠性感知 QoS 约束列为框架的直接扩展方向。"
        ],
        "placeholder": "Insert the marked-manuscript excerpts following Eq. (6) in Section III-C and the reliability-aware extension stated in the Conclusion (marked manuscript, pp. 5 and 14).",
        "evidence": "reliability",
    },
    {
        "section": "Reviewer 2",
        "number": "3",
        "title": "PSP/NSGA-II hyperparameters",
        "comment_en": (
            "The proposed PSP algorithm relies on NSGA-II for multi-objective optimization. Nevertheless, the manuscript lacks a detailed explanation of several key hyperparameters, such as population size {{Nsym}}, mutation probability, crossover probability, and maximum generation number {{Gsym}}."
        ),
        "comment_zh": "PSP 依赖 NSGA-II 进行多目标优化，但稿件缺少对关键超参数的详细说明，例如种群规模 {{Nsym}}、变异概率、交叉概率和最大迭代代数 {{Gsym}}。",
        "response_en": [
            "A complete parameter table has been added. The experiments use {{N50}}, {{G200}}, simulated binary crossover with {{pc_eta}}, polynomial mutation with {{pm_eta}}, hybrid-score weights {{alpha_beta}}, per-server service capacity {{Vj4}}, and anchor ratio {{rho05}}. The revised algorithm text also states the non-dominated rank and crowding-distance selection rule and the explicit capacity-repair procedure."
        ],
        "response_zh": [
            "修改稿增加了完整参数表。实验采用 {{N50}}、{{G200}}；模拟二进制交叉参数为 {{pc_eta}}；多项式变异参数为 {{pm_eta}}；混合评分权重为 {{alpha_beta}}；单服务器服务容量为 {{Vj4}}；锚点比例为 {{rho05}}。算法正文同时说明了非支配等级、拥挤距离选择规则以及显式容量修复流程。"
        ],
        "placeholder": "Insert the marked-manuscript excerpt containing Table II and the associated PSP parameter paragraph in Section V-A (marked manuscript, pp. 11-12).",
        "evidence": "parameters",
    },
    {
        "section": "Reviewer 2",
        "number": "4",
        "title": "Definition and normalization of Q",
        "comment_en": "In Fig. 5 and Fig. 7, the performance metric {{Qsym}} (normalized) is presented, but its exact normalization process and mathematical definition are not sufficiently described.",
        "comment_zh": "Fig. 5 和 Fig. 7 展示了 {{Qsym}} (normalized)，但其确切归一化过程和数学定义描述不充分。",
        "response_en": [
            "The revised manuscript defines min-max normalization over the pooled solutions of all compared methods within each server/user configuration using {{q_norm_C}} and {{q_norm_D}}. The scalar evaluation score is {{q_score}}. Lower {{Qsym}} denotes a better balanced cost-delay solution. The text further clarifies that {{Qsym}} is applied only after multi-objective optimization for scalar comparison and is not an objective used to generate the Pareto population."
        ],
        "response_zh": [
            "修改稿明确规定：在每个服务器/用户配置内，将全部对比方法的解合并后进行 min-max 归一化，采用 {{q_norm_C}} 和 {{q_norm_D}}；随后按 {{q_score}} 计算标量评价分数。{{Qsym}} 越低表示成本与时延的平衡解越好。正文还说明 {{Qsym}} 只在多目标优化之后用于标量比较，并不是生成 Pareto 种群的优化目标。"
        ],
        "placeholder": "Insert the marked-manuscript excerpt containing the Q normalization equations and interpretation in Section V-A (marked manuscript, p. 11).",
        "evidence": "q",
    },
    {
        "section": "Reviewer 2",
        "number": "5",
        "title": "Rationale for selecting the deterministic anchor size",
        "comment_en": (
            "The proposed hybrid initialization mechanism in Algorithm 2 introduces the parameter {{varpi_j}} to control the deterministic anchor size. However, the rationale behind selecting {{varpi_j}} is unclear."
        ),
        "comment_zh": "Algorithm 2 的混合初始化机制引入 {{varpi_j}} 控制确定性锚点大小，但稿件没有解释选择 {{varpi_j}} 的依据。",
        "response_en": [
            "The deterministic anchor is now defined proportionally as {{varpi_rule}}. This rule assigns half of each server's capacity to the highest-ranked deterministic services and reserves the remaining half for stochastic selection from lower-ranked candidates. It therefore provides an interpretable exploitation-exploration balance and scales automatically with heterogeneous server capacities. In the reported experiments, {{Vj4}} gives {{varpi2}}; a capacity of 8 gives {{varpi4}}."
        ],
        "response_zh": [
            "确定性锚点现按比例定义为 {{varpi_rule}}。该规则将每台服务器一半容量用于保留最高评分的确定性服务，另一半容量用于从较低排名候选中随机选择，从而形成可解释的利用-探索折中，并可随异构服务器容量自动缩放。本文实验中 {{Vj4}} 对应 {{varpi2}}；当容量为 8 时，对应 {{varpi4}}。"
        ],
        "placeholder": "Insert the marked-manuscript excerpts defining varpi_j in the hybrid-initialization subsection and Table II (marked manuscript, pp. 9-11).",
        "evidence": "varpi",
    },
    {
        "section": "Reviewer 2",
        "number": "6",
        "title": "Generalization to different geographical distributions and larger MEC settings",
        "comment_en": (
            "In Section V, all experiments are conducted using a dataset collected within approximately a 9 km region around Xizhimen Subway Station in Beijing. "
            "However, the manuscript does not discuss the generalization capability of the proposed framework under different geographical distributions, heterogeneous traffic densities, or larger-scale MEC environments."
        ),
        "comment_zh": "Section V 的实验均使用北京西直门地铁站周边约 9 km 区域的数据，但稿件没有讨论框架在不同地理分布、异构流量密度或更大规模 MEC 环境下的泛化能力。",
        "response_en": [
            "We conducted an additional complete two-stage experiment in a different real Beijing region using a pool of 2,215 deduplicated base-station coordinates. The evaluated instance contains 40 candidate base stations, 10 deployed servers, 130 users, and 8 service types. It changes the real base-station topology, doubles the candidate set relative to the primary experiment, and exhibits a more heterogeneous coverage-density structure.",
            "In Stage I, CLS obtained a cost of 2,304.7670, compared with 6,150.5741 for the best non-CLS initialization-only result, a 62.53% reduction. In Stage II, the three-seed means for PSP were {{HV10116}}, {{IGD0129}}, and Best {{Q02678}}. PSP achieved the best HV and IGD in each seed and the best mean values of all three metrics. The DQN baseline obtained mean Best {{Q05517}}, whereas PSP obtained 0.2678, corresponding to a 51.45% reduction. These results demonstrate that the complete framework remains effective under a different real deployment topology, a larger candidate set, and heterogeneous spatial demand."
        ],
        "response_zh": [
            "我们使用 2,215 个去重后的北京真实基站坐标，在不同真实区域完成了一组从 Stage I 到 Stage II 的完整实验。测试实例包含 40 个候选基站、10 台部署服务器、130 个用户和 8 类服务。相对于主实验，该实例改变了真实基站拓扑，将候选基站数量扩大一倍，并呈现更异构的覆盖密度结构。",
            "Stage I 中，CLS 的成本为 2,304.7670，最佳非 CLS 初始化结果为 6,150.5741，下降 62.53%。Stage II 的三种子平均结果中，PSP 的结果为 {{HV10116}}、{{IGD0129}} 和 Best {{Q02678}}；PSP 在每个种子下都取得最佳 HV 和 IGD，并取得三项指标的最佳均值。DQN 的平均结果为 Best {{Q05517}}，PSP 为 0.2678，下降 51.45%。结果表明，完整框架在不同真实部署拓扑、更大的候选集合和异构空间需求下仍保持有效。"
        ],
        "evidence": "generalization",
    },
    {
        "section": "Reviewer 2",
        "number": "7",
        "title": "Learning-based service-placement baseline",
        "comment_en": (
            "The comparison baselines in Stage II mainly include heuristic initialization strategies (GCP, GDP) and standard NSGA-II initialization. "
            "However, the manuscript does not compare against recent learning-based service placement methods, such as deep reinforcement learning or graph neural network based approaches."
        ),
        "comment_zh": "Stage II 的基线主要包括 GCP、GDP 等启发式初始化和标准 NSGA-II 初始化，但没有与近期基于学习的服务放置方法比较，例如深度强化学习或图神经网络方法。",
        "response_en": [
            "We included a Deep Q-Network provisioning baseline. Stage-II provisioning is represented as a sequence of server-slot decisions in which the Q-network selects a service type or an empty action from demand, deployment-cost, current-selection, and cost-delay preference information. The resulting service matrix is evaluated by exactly the same cost and delay functions as the other methods. The implementation uses one hidden layer with 64 units and trains for 320 episodes for each preference weight in {{preference_weights}}; the learning rate is {{learning_rate}}, the discount factor is 0.98, the mini-batch size is 64, and the replay capacity is 12,000.",
            "Across both primary Stage-II experiment series, PSP attained the lowest normalized {{Qsym}} in every reported case. In the representative 10-server/130-user setting, PSP achieved {{Q03282}}, compared with {{Q06125}} for DQN. Because DQN returns preference-conditioned deployment points rather than an evolutionary population, it is compared through the common scalar {{Qsym}} measure and is not presented as a Pareto curve."
        ],
        "response_zh": [
            "我们引入了 Deep Q-Network 服务配置基线。Stage II 被表示为一系列服务器槽位决策，Q 网络依据需求、配置成本、当前已选服务和成本-时延偏好，为每个槽位选择一种服务或空动作。输出的服务配置矩阵采用与其他方法完全相同的成本和时延函数进行评价。实现使用一个含 64 个单元的隐藏层，并对每个偏好权重 {{preference_weights}} 训练 320 个 episode；学习率为 {{learning_rate}}，折扣因子 0.98，mini-batch 为 64，经验回放容量为 12,000。",
            "在两组主要 Stage II 实验中，PSP 在所有展示配置下都取得最低归一化 {{Qsym}}。以 10 台服务器/130 个用户为例，PSP 的结果为 {{Q03282}}，DQN 为 {{Q06125}}。由于 DQN 返回的是偏好条件化部署点，而不是进化算法种群，因此通过统一标量指标 {{Qsym}} 进行比较，不将其画成 Pareto 曲线。"
        ],
        "placeholder": "Insert the marked-manuscript excerpts containing the concise DQN baseline description, the two Stage-II comparison figures, and the accompanying result paragraph in Section V (marked manuscript, pp. 12-13).",
        "evidence": "dqn",
    },
    {
        "section": "Reviewer 2",
        "number": "8",
        "title": "Quantitative Pareto metrics",
        "comment_en": (
            "In Fig. 8, the Pareto fronts of different algorithms are illustrated, but no quantitative Pareto evaluation metrics (e.g., Hypervolume, IGD, Spacing, or Spread) are provided. "
            "Relying only on visual comparison may not be sufficiently rigorous."
        ),
        "comment_zh": "Fig. 8 展示了不同算法的 Pareto 前沿，但没有提供 Hypervolume、IGD、Spacing 或 Spread 等定量 Pareto 指标。仅依靠视觉比较可能不够严谨。",
        "response_en": [
            "We added Hypervolume (HV) and Inverted Generational Distance (IGD), together with Best {{Qsym}}, for the representative 10-server/130-user configuration. HV measures dominated objective-space volume relative to the common reference point {{ref_point}} and is maximized; IGD measures the mean distance from the common non-dominated reference front to a method's front and is minimized. PSP achieved {{HV09470}}, {{IGD00016}}, and Best {{Q03282}}, outperforming NS-P, GCP, and GDP on all three measures.",
            "DQN provides a limited set of preference-conditioned solutions rather than an equal-cardinality population-based Pareto front. It is therefore evaluated through the common Best {{Qsym}} measure, while HV and IGD are reported for the four population-based strategies."
        ],
        "response_zh": [
            "我们在代表性的 10 台服务器/130 个用户配置下增加了 Hypervolume (HV)、Inverted Generational Distance (IGD) 和 Best {{Qsym}}。HV 衡量相对于公共参考点 {{ref_point}} 的支配目标空间体积，越高越好；IGD 衡量公共非支配参考前沿到方法前沿的平均距离，越低越好。PSP 的结果为 {{HV09470}}、{{IGD00016}} 和 Best {{Q03282}}，三项均优于 NS-P、GCP 和 GDP。",
            "DQN 输出有限个偏好条件化解，不是等规模的种群型 Pareto 前沿，因此使用统一 Best {{Qsym}} 比较；HV 和 IGD 则用于四种种群型策略。"
        ],
        "placeholder": "Insert the marked-manuscript excerpts containing the HV/IGD definitions, the quantitative metric table, and Fig. 8 in Section V (marked manuscript, pp. 11-13).",
        "evidence": "pareto",
    },
    {
        "section": "Reviewer 2",
        "number": "9",
        "title": "Suggested references",
        "comment_en": (
            "The following works are closely related to MEC and server placement, and thus should not be overlooked. "
            "[1] Latency-Aware Service Deployment and Peer Offloading: A Long-Term Optimization Framework for Satellite Edge Computing. "
            "[2] Latency-Aware Task Offloading in Multi-Tier SAGIN With FSO-Enabled Mobile Edge Computing. "
            "[3] Novel Breakout Local Search for Offloading Tasks in Multi-Tiered Cloud Environment Considering Transmission and Processing. "
            "[4] Long-Term Max-Min Fairness Guarantee Mechanism for Integrated Multi-RAT and MEC Networks. "
            "[5] Dynamic Energy Cost Conservation for Distributed Edge Clouds Utilizing Online Mini-Batch Learning. "
            "[6] Mobile Edge Computing Offloading for Static Users in a Free Space Optical Communications-Enabled Satellite-Air-Ground Integrated Network."
        ),
        "comment_zh": "审稿人列出六篇与 MEC、服务器放置、任务卸载、长期资源管理和在线学习相关的工作，并建议不要忽略。",
        "response_en": [
            "Following the Associate Editor's guidance, we evaluated the six works individually and incorporated four studies with direct methodological relevance: Feng et al. on latency-aware service deployment and peer offloading; Kato et al. on breakout local search for transmission- and processing-aware offloading; Jing et al. on long-term max-min fairness in integrated multi-RAT/MEC networks; and Jing et al. on online mini-batch learning for dynamic energy-cost conservation in distributed edge clouds. The Related Work section now explains how these studies complement the present focus on interpretable server planning and capacity-constrained multi-objective service provisioning.",
            "The two remaining studies specifically address task offloading in FSO-enabled satellite-air-ground integrated networks. Their network architecture, optical-link assumptions, and decision variables differ substantially from the terrestrial MEC service-provisioning model considered here; consequently, they were not included solely to increase the citation count."
        ],
        "response_zh": [
            "根据副编辑的指导，我们逐篇评估六篇文献，并引入其中四篇与方法直接相关的研究：Feng 等人的时延感知服务部署与对等卸载；Kato 等人的传输/处理感知突破局部搜索卸载；Jing 等人的集成多 RAT/MEC 长期最大最小公平性；以及另一项 Jing 等人关于分布式边缘云动态能源成本的在线小批量学习。Related Work 现已说明这些研究如何从动态卸载、局部搜索、公平性和学习型能源管理角度补充本文的可解释服务器规划与容量受限多目标服务配置。",
            "其余两篇专门研究 FSO 支持的空天地一体化网络任务卸载，其网络架构、光链路假设和决策变量与本文地面 MEC 服务配置模型差异较大，因此没有仅为增加引用数量而纳入。"
        ],
        "placeholder": "Insert the marked-manuscript excerpts containing the new Related Work paragraph and references [33]-[36] (marked manuscript, pp. 3 and 14).",
        "evidence": "references",
    },
]


REVISED_EXCERPTS = {
    "fig1": {
        "location": "Section I-A, Fig. 1 and the accompanying discussion",
        "blocks": [
            {
                "type": "caption",
                "en": "Fig. 1. MEC system model and service-provisioning scenario.",
                "zh": "图 1. MEC 系统模型与服务配置场景。",
            },
            {
                "type": "paragraph",
                "en": "Server deployment and service provisioning have a significant impact on the profit of mobile network operators. Figure 1 illustrates a scenario with five base stations. The solid black contours partition the topology into interconnected service regions, the black dashed circles denote the coverage areas of individual base stations, and the red dotted circles identify base stations equipped with edge servers. The blue bidirectional arrows represent cloud-edge communication, while the red dashed arrows indicate inter-region request forwarding. Distinct colors represent service types {{service_range}}, and each user color indicates the corresponding requested service. The service blocks beside an edge server show its instantiated services; the question-mark slot in Area 3 denotes a provisioning decision under limited server capacity.",
                "zh": "服务器部署和服务配置对移动网络运营商的收益具有显著影响。图 1 展示了一个包含五个基站的场景。黑色实线轮廓将拓扑划分为相互连接的服务区域，黑色虚线圆表示各基站的覆盖区域，红色点线圆标识配置了边缘服务器的基站。蓝色双向箭头表示云边通信，红色虚线箭头表示跨区域请求转发。不同颜色表示服务类型 {{service_range}}，每个用户的颜色表示其对应的请求服务。边缘服务器旁的服务色块表示其实例化的服务；Area 3 中带问号的槽位表示服务器容量受限条件下尚待确定的配置决策。",
            },
            {
                "type": "paragraph",
                "en": "Each service replica serves a user request and generates income. Because an edge server can host only a subset of the service catalog, a request without a local replica is forwarded to another server that provides the requested service, incurring additional transmission cost. For example, the server in Area 4 provisions services {{s_3}} and {{s_6}}; requests for the other service types require inter-area forwarding.",
                "zh": "每个服务副本处理一个用户请求并产生收益。由于边缘服务器只能承载服务目录中的一部分服务，若本地没有相应副本，请求将被转发到能够提供该服务的其他服务器，从而产生额外传输成本。例如，Area 4 的服务器配置了服务 {{s_3}} 和 {{s_6}}；对其他服务类型的请求需要跨区域转发。",
            },
        ],
    },
    "notation": {
        "location": "Section III-A and Section IV-B",
        "blocks": [
            {
                "type": "paragraph",
                "en": "Equation (7) defines the server-deployment variable {{y_jk}}, service-provisioning variable {{z_jw}}, and user-service association variable {{a_ijw}}.",
                "zh": "公式 (7) 分别定义服务器部署变量 {{y_jk}}、服务配置变量 {{z_jw}} 和用户-服务关联变量 {{a_ijw}}。",
            },
            {
                "type": "paragraph",
                "en": "A service-provisioning scheme is encoded by a binary matrix {{indiv_matrix}}. Each row corresponds to a deployed server and each column to a service type. The entry {{indiv_jw}} is the encoded counterpart of {{z_jw}}: it equals 1 when service {{s_w}} is provisioned on server {{m_j}}, and 0 otherwise.",
                "zh": "服务配置方案编码为二进制矩阵 {{indiv_matrix}}。每一行对应一台已部署服务器，每一列对应一种服务类型。矩阵项 {{indiv_jw}} 是 {{z_jw}} 的编码对应项：当服务 {{s_w}} 配置在服务器 {{m_j}} 上时取 1，否则取 0。",
            },
        ],
    },
    "joint": {
        "location": "Section III-D, decomposition rationale",
        "blocks": [
            {
                "type": "paragraph",
                "en": "The decomposition follows the operational hierarchy of MEC planning. Server deployment is a long-term infrastructure decision governed by budget, coverage, and geographic demand, whereas service instances are provisioned after the physical server locations are known. Stage 1 therefore resolves the spatial variables that determine deployment and transmission costs, and Stage 2 optimizes service instances and user associations on the selected infrastructure. Passing the selected locations, user assignments, and Stage-I cost to Stage 2 preserves the cost-latency dependence while substantially reducing the decision space of the service-provisioning search.",
                "zh": "该分解遵循 MEC 规划的实际决策层级。服务器部署是受预算、覆盖范围和地理需求约束的长期基础设施决策，而服务实例则在物理服务器位置确定后进行配置。因此，Stage 1 求解决定部署成本和传输成本的空间变量，Stage 2 在选定基础设施上优化服务实例和用户关联。将选定位置、用户分配和 Stage-I 成本传递至 Stage 2，可在大幅缩小服务配置搜索决策空间的同时保留成本-时延依赖关系。",
            }
        ],
    },
    "fig2": {
        "location": "Section IV and Algorithm 3",
        "blocks": [
            {
                "type": "caption",
                "en": "Fig. 2. Overall architecture of MOS².",
                "zh": "图 2. MOS² 的总体架构。",
            },
            {
                "type": "paragraph",
                "en": "Given the fixed server locations, PSP constructs a hybrid initial population and applies NSGA-II to optimize provisioning cost and access delay. Crossover and mutation produce each offspring, whose entries are first rounded to binary values. For every server {{m_j}}, PSP then checks {{capacity_z}}. If the capacity is exceeded, selected service entries are randomly deactivated until feasibility is restored. Only capacity-feasible offspring are evaluated and passed to non-dominated sorting; thus, feasibility is enforced by repair rather than by a penalty function. The evolutionary process returns a set of provisioning schemes representing different cost-delay trade-offs.",
                "zh": "在服务器位置固定后，PSP 构造混合初始种群，并应用 NSGA-II 优化服务配置成本和访问时延。交叉与变异产生子代后，首先将其各编码项舍入为二进制值。随后，PSP 对每台服务器 {{m_j}} 检查 {{capacity_z}}。若容量超限，则随机停用已选服务项，直至恢复可行性。只有满足容量约束的子代才进入目标评估和非支配排序；因此，可行性通过修复而非惩罚函数来保证。进化过程最终返回一组表示不同成本-时延权衡的服务配置方案。",
            },
            {
                "type": "paragraph",
                "en": "After crossover and mutation, every encoded entry is rounded to 0 or 1. If {{capacity_indiv_exceed}} for any server {{m_j}}, selected entries in that row are randomly set to 0 until the capacity constraint is satisfied. Objective evaluation and non-dominated sorting are then performed on the repaired feasible individual.",
                "zh": "交叉和变异后，每个编码项都被舍入为 0 或 1。若任一服务器 {{m_j}} 满足 {{capacity_indiv_exceed}}，则随机将该行中的已选项置为 0，直至满足容量约束；随后才对修复后的可行个体进行目标评估和非支配排序。",
            },
        ],
    },
    "proofreading": {
        "location": "Sections I and III, representative corrected passages",
        "blocks": [
            {
                "type": "paragraph",
                "en": "In this paper, we consider a hierarchical network architecture comprising the cloud data center, edge servers, and mobile end users, as illustrated in Figure 1.",
                "zh": "本文考虑由云数据中心、边缘服务器和移动终端用户构成的分层网络架构，如图 1 所示。",
            },
            {
                "type": "paragraph",
                "en": "At the user layer, the set of mobile users is denoted by {{user_set}}, where {{u_i}} represents the i-th user.",
                "zh": "在用户层，移动用户集合记为 {{user_set}}，其中 {{u_i}} 表示第 i 个用户。",
            },
        ],
    },
    "cls": {
        "location": "Section V-B, Stage-I initialization-sensitivity analysis",
        "blocks": [
            {
                "type": "paragraph",
                "en": "To examine the effect of the initial deployment set in Algorithm 1, we compare Random, Density, DistSum, marginal Greedy, and Diverse initialization over 50 runs. For each configuration, the reported gap is the percentage difference between the final cost and the best final cost observed under the same server/user setting. As shown in Fig. 9, the five strategies reach the same best final cost in nearly all fixed-130-user cases; the only nonzero entry is the 2.10% mean gap of Random at 10 servers. Under the 10-server/150-user setting, Random obtains a 1.27% mean gap, whereas marginal Greedy reaches 15.88%. These results show that CLS is generally stable across initializations and that random initialization avoids a systematic preference for a poorer local optimum.",
                "zh": "为考察 Algorithm 1 中初始部署集合的影响，我们在 50 次运行中比较 Random、Density、DistSum、marginal Greedy 和 Diverse 五种初始化。对于每种配置，所报告的 gap 是最终成本相对于同一服务器/用户设置下观测到的最佳最终成本的百分比差值。如 Fig. 9 所示，在固定 130 个用户的几乎所有情形中，五种策略均达到相同的最佳最终成本；唯一的非零项是 10 台服务器时 Random 的 2.10% 平均 gap。在 10 台服务器/150 个用户的设置下，Random 的平均 gap 为 1.27%，而 marginal Greedy 达到 15.88%。这些结果表明 CLS 对初始化总体稳定，且随机初始化可避免对较差局部最优形成系统性偏好。",
            },
            {
                "type": "caption",
                "en": "Fig. 9. Initialization sensitivity of CLS under two server/user settings.",
                "zh": "图 9. 两种服务器/用户设置下 CLS 的初始化敏感性。",
            },
        ],
    },
    "reliability": {
        "location": "Section III-C and Conclusion",
        "blocks": [
            {
                "type": "paragraph",
                "en": "Equation (6) provides latency-oriented QoS by requiring the end-to-end latency of a served request to remain below the user's tolerable bound {{tilde_D_i}}. For reliability-critical applications, latency compliance can be complemented by packet-loss, link-availability, and service-interruption constraints, preventing a low-latency solution from being selected when service continuity is insufficient.",
                "zh": "公式 (6) 通过要求已服务请求的端到端时延不超过用户可容忍上界 {{tilde_D_i}}，提供面向时延的 QoS 保障。对于可靠性关键型应用，可在时延约束之外补充丢包率、链路可用性和服务中断约束，避免在服务连续性不足时选取低时延方案。",
            },
            {
                "type": "paragraph",
                "en": "Future work will extend the framework with reliability-aware QoS constraints for packet loss, link availability, and service interruption, together with learning-based demand prediction and multi-domain coordination.",
                "zh": "未来工作将通过面向丢包、链路可用性和服务中断的可靠性感知 QoS 约束扩展该框架，并进一步研究基于学习的需求预测和多域协同。",
            },
        ],
    },
    "parameters": {
        "location": "Section V-A and Table II",
        "blocks": [
            {
                "type": "paragraph",
                "en": "The evolutionary settings used for PSP are summarized in Table II. The deterministic anchor size is defined proportionally as {{varpi_rule}}. Hence, for the experimental capacity {{Vj4}}, two slots retain the highest-scoring services and two slots preserve stochastic exploration. This equal allocation provides an interpretable exploitation-exploration balance and applies directly to servers with different capacities.",
                "zh": "PSP 使用的进化参数汇总于 Table II。确定性锚点大小按比例定义为 {{varpi_rule}}。因此，在实验容量 {{Vj4}} 下，两个槽位保留评分最高的服务，另外两个槽位保留随机探索。该等比例分配形成可解释的利用-探索平衡，并可直接适用于不同容量的服务器。",
            }
        ],
    },
    "q": {
        "location": "Section V-A, normalization and scalar evaluation",
        "blocks": [
            {
                "type": "paragraph",
                "en": "For each server/user configuration, cost and delay are normalized over the pooled solutions of the compared methods as",
                "zh": "对于每种服务器/用户配置，在所有对比方法的合并解集上对成本和时延进行如下归一化：",
            },
            {"type": "equation", "key": "q_norm_C", "zh": "成本的 min-max 归一化公式。"},
            {"type": "equation", "key": "q_norm_D", "zh": "时延的 min-max 归一化公式。"},
            {
                "type": "paragraph",
                "en": "and the scalar evaluation score is",
                "zh": "随后，标量评价分数定义为：",
            },
            {"type": "equation", "key": "q_score", "zh": "成本与时延等权的标量评价分数。"},
            {
                "type": "paragraph",
                "en": "Lower {{Qsym}} indicates a better balanced solution. Hypervolume (HV) measures the dominated objective-space volume relative to the reference point {{ref_point}} and is maximized, whereas inverted generational distance (IGD) measures the mean distance from the common non-dominated reference front to a method's front and is minimized.",
                "zh": "{{Qsym}} 越低表示成本与时延的平衡解越好。Hypervolume (HV) 衡量相对于参考点 {{ref_point}} 的支配目标空间体积，越大越好；Inverted Generational Distance (IGD) 衡量公共非支配参考前沿到某方法前沿的平均距离，越小越好。",
            },
        ],
    },
    "varpi": {
        "location": "Section IV-B and Section V-A",
        "blocks": [
            {
                "type": "paragraph",
                "en": "PSP uses a hybrid initialization before the NSGA-II evolutionary loop. For each server, it constructs one ranking that favors low provisioning cost and another that favors services requested frequently by locally assigned users. The two rankings are fused by a hybrid score. High-scoring services form a deterministic anchor set, and the remaining capacity is filled by stochastic sampling from lower-ranked candidates. The resulting base configuration combines exploitation of high-quality candidates with exploration of alternative service combinations.",
                "zh": "PSP 在 NSGA-II 进化循环之前采用混合初始化。对于每台服务器，分别构造偏向低配置成本的排序和偏向本地分配用户高频请求服务的排序，并通过混合评分融合。高评分服务构成确定性锚点集合，剩余容量则从低排名候选中随机采样填充。由此得到的基础配置同时利用高质量候选并探索替代服务组合。",
            },
            {
                "type": "paragraph",
                "en": "Let {{alpha_beta_range}} balance provisioning cost and local demand. Let {{varpi_j}} ({{varpi_bounds}}) denote the deterministic anchor size on edge server {{m_j}}. The hybrid score of candidate service {{s_w}} is",
                "zh": "令 {{alpha_beta_range}} 用于平衡配置成本与本地需求。令 {{varpi_j}}（{{varpi_bounds}}）表示边缘服务器 {{m_j}} 上的确定性锚点大小。候选服务 {{s_w}} 的混合评分为：",
            },
            {"type": "equation", "key": "hybrid_score", "zh": "候选服务的混合评分公式。"},
            {
                "type": "paragraph",
                "en": "where {{w_cost_j}} and {{w_req_j}} are descending rank scores derived from provisioning cost and the request frequency of users assigned to {{m_j}}, respectively. We set {{varpi_rule}}, reserving half of the capacity for high-ranked deterministic services and half for stochastic selection from the remaining candidates.",
                "zh": "其中，{{w_cost_j}} 和 {{w_req_j}} 分别是由配置成本及分配到 {{m_j}} 的用户请求频率得到的降序排名分数。本文设置 {{varpi_rule}}，将一半容量用于高排名确定性服务，另一半用于从其余候选中随机选择。",
            },
        ],
    },
    "dqn": {
        "location": "Section V-A and Section V-B",
        "blocks": [
            {
                "type": "paragraph",
                "en": "Deep Q-Network Provision (DQN): A Q-network sequentially selects a service or an empty action for each server slot using demand, deployment cost, and a cost-delay preference. The resulting deployment is evaluated by the same objectives as the other methods.",
                "zh": "Deep Q-Network Provision (DQN)：Q 网络根据需求、部署成本和成本-时延偏好，依次为每个服务器槽位选择一种服务或空动作；所得部署方案使用与其他方法相同的目标函数进行评价。",
            },
            {
                "type": "paragraph",
                "en": "We evaluate Stage II under two complementary configurations: (i) 10 deployed servers with 100, 130, 150, and 180 users, and (ii) 130 users with 5, 10, 15, and 20 deployed servers. Figures 5 and 7 report the minimum normalized score {{Qsym}} obtained by each method. PSP achieves the lowest {{Qsym}} in every reported case across both experimental series. Its hybrid initialization consistently improves upon random NS-P and the single-criterion GCP and GDP initializations; PSP also achieves lower {{Qsym}} than DQN across all tested scales.",
                "zh": "我们在两组互补配置下评估 Stage II：(i) 固定部署 10 台服务器，用户数为 100、130、150 和 180；(ii) 固定 130 个用户，部署服务器数为 5、10、15 和 20。Figures 5 和 7 给出各方法获得的最小归一化分数 {{Qsym}}。在两组实验的全部配置中，PSP 均取得最低 {{Qsym}}。其混合初始化始终优于随机 NS-P 以及单准则 GCP 和 GDP 初始化；在所有测试规模下，PSP 的 {{Qsym}} 也均低于 DQN。",
            },
        ],
    },
    "pareto": {
        "location": "Section V-A and Section V-B",
        "blocks": [
            {
                "type": "paragraph",
                "en": "Lower {{Qsym}} indicates a better balanced solution. Hypervolume (HV) measures the dominated objective-space volume relative to the reference point {{ref_point}} and is maximized, whereas inverted generational distance (IGD) measures the mean distance from the common non-dominated reference front to a method's front and is minimized.",
                "zh": "{{Qsym}} 越低表示成本与时延的平衡解越好。Hypervolume (HV) 衡量相对于参考点 {{ref_point}} 的支配目标空间体积，越大越好；Inverted Generational Distance (IGD) 衡量公共非支配参考前沿到某方法前沿的平均距离，越小越好。",
            },
            {
                "type": "paragraph",
                "en": "Figure 8 compares the cost-delay fronts generated by the four NSGA-II-based provisioning strategies with 130 users. Table III gives the numerical results for the 10-server setting. PSP attains the largest HV (0.9470), the smallest IGD (0.0016), and the smallest {{Qsym}} (0.3282), confirming both broader objective-space coverage and closer convergence to the common reference front. DQN produces a limited set of scalarized solutions rather than a population-based Pareto front; it is therefore compared through {{Qsym}} and is not included in the HV/IGD comparison.",
                "zh": "Figure 8 比较了 130 个用户下四种基于 NSGA-II 的服务配置策略所生成的成本-时延前沿。Table III 给出 10 台服务器设置下的数值结果。PSP 取得最大 HV (0.9470)、最小 IGD (0.0016) 和最小 {{Qsym}} (0.3282)，表明其目标空间覆盖更广且更接近公共参考前沿。DQN 产生的是有限个标量化解，而不是种群型 Pareto 前沿，因此仅通过 {{Qsym}} 比较，不纳入 HV/IGD 对比。",
            },
        ],
    },
    "references": {
        "location": "Section II-C, Related Work",
        "blocks": [
            {
                "type": "paragraph",
                "en": "Optimization and learning methods also address complementary MEC resource-management objectives. Feng et al. [33] jointly optimized latency-aware service deployment and peer offloading over multiple timescales. Kato et al. [34] applied breakout local search to transmission- and processing-aware task offloading in a multi-tier cloud environment. Jing et al. [35] studied long-term max-min fairness for task splitting and resource allocation in integrated multi-RAT/MEC networks, while Jing et al. [36] used online mini-batch learning for dynamic energy-cost conservation in distributed edge clouds. These studies address dynamic offloading, fairness, or energy management; the present work focuses on interpretable server planning and capacity-constrained multi-objective service provisioning under a common cost-latency model.",
                "zh": "优化与学习方法也从互补角度研究 MEC 资源管理目标。Feng 等人 [33] 在多时间尺度下联合优化时延感知服务部署和对等卸载。Kato 等人 [34] 将突破局部搜索用于多层云环境中同时考虑传输和处理的任务卸载。Jing 等人 [35] 研究集成多 RAT/MEC 网络中任务拆分和资源分配的长期最大最小公平性；另一项 Jing 等人的工作 [36] 使用在线小批量学习实现分布式边缘云的动态能源成本节约。这些研究侧重动态卸载、公平性或能源管理；本文则聚焦共同成本-时延模型下可解释的服务器规划与容量受限多目标服务配置。",
            }
        ],
    },
}


def add_evidence(doc, key):
    if key == "fig1":
        add_figure(doc, ASSETS / "fig1_revised.png", "Response Fig. R1-1. Revised system model and service-provisioning scenario.", width=3.85)
    elif key == "notation":
        add_data_table(
            doc,
            ["Decision", "Revised variable", "Definition"],
            [
                ["Server deployment", "{{y_jk}}", "Server {{m_j}} is deployed at base station {{b_k}}"],
                ["Service provisioning", "{{z_jw}}", "Service {{s_w}} is provisioned on server {{m_j}}"],
                ["User association", "{{a_ijw}}", "User {{u_i}} is served by {{s_w}} on {{m_j}}"],
            ],
            widths=[1.45, 1.25, 3.75],
            caption="Notation revision summary",
        )
    elif key == "joint":
        add_figure(doc, ASSETS / "joint_exact_vs_mos2_seed42.png", "Response Fig. R1-2. Representative small-scale Pareto comparison between Joint-Exact and MOS\u00b2-PSP (seed 42; response-only evidence).", width=5.25)
        add_data_table(
            doc,
            ["Seed", "HV gap", "IGD", "Best Q gap", "Runtime Joint/MOS\u00b2"],
            joint_rows(),
            widths=[0.75, 1.05, 0.95, 1.15, 1.85],
            bold_cells={(3, 0), (3, 1), (3, 2), (3, 3), (3, 4)},
            caption="Small-scale exact joint-optimization results",
            font_size=8.7,
        )
    elif key == "fig2":
        add_figure(doc, ASSETS / "fig2_evolution_panel.png", "Response Fig. R1-3. Revised Evolutionary Optimization flow showing binary rounding, capacity checking, repair, evaluation, and population update.", width=5.5)
    elif key == "readability":
        add_figure(doc, ASSETS / "stage1_scale.png", "Response Fig. R1-4. Re-exported Stage-I multi-panel result.", width=5.65)
        add_figure(doc, ASSETS / "stage2_fixed_users_dqn.png", "Response Fig. R1-5. Re-exported Stage-II multi-panel result at double-column width.", width=5.95)
    elif key == "proofreading":
        add_data_table(
            doc,
            ["Original", "Revised"],
            [
                ["s2 , s3 and s5", "s2, s3, and s5"],
                ["U the set of", "U is the set of"],
                ["4rd / 5rd", "4th / 5th"],
                ["hierarchy network architecture", "hierarchical network architecture"],
            ],
            widths=[3.2, 3.2],
            caption="Representative proofreading corrections",
        )
    elif key == "cls":
        add_figure(doc, ASSETS / "cls_initialization_sensitivity.png", "Response Fig. R2-1. CLS initialization sensitivity across the fixed-130-user configurations and the 10-server/150-user stress case.", width=5.85)
        add_data_table(
            doc,
            ["Configuration", "Random mean gap", "Marginal-greedy gap", "Other initializers"],
            [
                ["5/130", "0.00%", "0.00%", "0.00%"],
                ["10/130", "2.10%", "0.00%", "0.00%"],
                ["15/130", "0.00%", "0.00%", "0.00%"],
                ["20/130", "0.00%", "0.00%", "0.00%"],
                ["10/150", "1.27%", "15.88%", "Not used in the focused pairwise panel"],
            ],
            widths=[1.1, 1.35, 1.5, 2.55],
            caption="Mean gap to the best final CLS cost over 50 runs",
            font_size=8.6,
        )
    elif key == "parameters":
        add_data_table(
            doc,
            ["Parameter", "Setting", "Role"],
            [
                ["Population size", "{{N50}}", "Number of individuals"],
                ["Maximum generations", "{{G200}}", "Evolutionary iterations"],
                ["SBX crossover", "{{pc_eta}}", "Crossover probability and index"],
                ["Polynomial mutation", "{{pm_eta}}", "Mutation probability and index"],
                ["Hybrid-score weights", "{{alpha_beta}}", "Cost-demand balance"],
                ["Service capacity", "{{Vj4}}", "Services per server"],
                ["Anchor ratio", "{{rho05}}", "Deterministic capacity share"],
            ],
            widths=[1.6, 2.15, 2.65],
            caption="PSP/NSGA-II parameter settings",
        )
    elif key == "q":
        pass
    elif key == "varpi":
        add_figure(doc, ASSETS / "hybrid_initialization.png", "Response Fig. R2-2. Hybrid initialization showing deterministic anchoring and stochastic filling.", width=4.9)
    elif key == "generalization":
        add_data_table(
            doc,
            ["Property", "Value"],
            [
                ["Real Beijing station pool", "2,215 deduplicated coordinates"],
                ["Candidate stations / deployed servers", "40 / 10"],
                ["Users / service types", "130 / 8"],
                ["Coverage-density coefficient of variation", "0.3359"],
            ],
            widths=[3.35, 3.05],
            caption="Response-only real-region experiment configuration",
        )
        add_figure(doc, ASSETS / "real_region_topology.png", "Response Fig. R2-3. Real-region topology, heterogeneous user requests, and CLS-selected edge-server locations.", width=5.25)
        add_data_table(
            doc,
            ["Stage-I method", "Objective (lower is better)"],
            [
                ["CLS", "2,304.7670"],
                ["Best random trial", "6,150.5741"],
                ["Random-trial mean", "13,540.5175"],
                ["Density", "22,206.4751"],
                ["Distance-Sum", "28,513.3397"],
                ["Marginal Greedy", "28,513.3397"],
                ["Density-Diverse", "19,031.7959"],
            ],
            widths=[3.2, 3.2],
            bold_cells={(0, 0), (0, 1)},
            caption="Stage-I results in the new real-region instance",
        )
        add_data_table(
            doc,
            ["Method", "HV mean \u00b1 std", "IGD mean \u00b1 std", "Best Q mean \u00b1 std"],
            real_region_stage2_rows(),
            widths=[1.0, 1.7, 1.7, 2.0],
            bold_cells={(3, 0), (3, 1), (3, 2), (3, 3)},
            caption="Stage-II results over seeds 42, 43, and 44",
            font_size=8.6,
        )
        add_figure(doc, ASSETS / "real_region_bestq.png", "Response Fig. R2-4. Mean Best Q in the real-region experiment; bars report the three-seed means.", width=5.25)
        add_data_table(
            doc,
            ["Method", "Best Q mean", "Std."],
            real_region_bestq_rows(),
            widths=[2.4, 2.0, 2.0],
            bold_cells={(3, 0), (3, 1), (3, 2)},
            caption="Balanced-solution comparison including DQN",
        )
    elif key == "dqn":
        add_figure(doc, ASSETS / "stage2_fixed_servers_dqn.png", "Response Fig. R2-5. Stage-II Best Q with 10 deployed servers and increasing user populations.", width=5.95)
        add_figure(doc, ASSETS / "stage2_fixed_users_dqn.png", "Response Fig. R2-6. Stage-II Best Q with 130 users and increasing deployed-server counts.", width=5.95)
        add_data_table(
            doc,
            ["Series", "Servers/users", "PSP Q", "DQN Q", "PSP reduction"],
            stage2_dqn_rows(),
            widths=[1.45, 1.15, 0.9, 0.9, 1.25],
            bold_cells={(i, 2) for i in range(8)},
            caption="Direct PSP-DQN comparison using the common normalized Q",
            font_size=8.4,
        )
    elif key == "pareto":
        add_figure(doc, ASSETS / "pareto_10_130.png", "Response Fig. R2-7. Pareto fronts for the representative 10-server/130-user configuration.", width=5.35)
        add_data_table(
            doc,
            ["Method", "HV (higher)", "IGD (lower)", "Best Q (lower)"],
            [
                ["NS-P", "0.8191", "0.0785", "0.3894"],
                ["GCP", "0.8596", "0.0492", "0.3550"],
                ["GDP", "0.8945", "0.0326", "0.3363"],
                ["PSP", "0.9470", "0.0016", "0.3282"],
                ["DQN", "N/A", "N/A", "0.6125"],
            ],
            widths=[1.35, 1.55, 1.55, 1.75],
            bold_cells={(3, 0), (3, 1), (3, 2), (3, 3)},
            caption="Quantitative Pareto and balanced-solution metrics",
        )
    elif key == "references":
        refs = [
            "C. Feng, M. Yang, Z. Jing, T. Q. S. Quek, and M. Mei, \"Latency-Aware Service Deployment and Peer Offloading: A Long-Term Optimization Framework for Satellite Edge Computing,\" IEEE Internet of Things Journal, vol. 13, no. 1, pp. 405-417, 2026.",
            "M. Kato, T. K. Rodrigues, and S. Verma, \"Novel Breakout Local Search for Offloading Tasks in Multi-Tiered Cloud Environment Considering Transmission and Processing,\" Proc. IEEE IC-NIDC, pp. 269-274, 2025.",
            "Z. Jing, Q. Yang, M. Qin, J. Li, and K. S. Kwak, \"Long-Term Max-Min Fairness Guarantee Mechanism for Integrated Multi-RAT and MEC Networks,\" IEEE Transactions on Vehicular Technology, vol. 70, no. 3, pp. 2478-2492, 2021.",
            "Z. Jing, X. Wang, Q. Yang, M. Mei, and Y. Wu, \"Dynamic Energy Cost Conservation for Distributed Edge Clouds Utilizing Online Mini-Batch Learning,\" Proc. IEEE PIMRC, pp. 1-6, 2023.",
        ]
        for i, ref in enumerate(refs, start=1):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.05)
            r = p.add_run(ref)
            set_run_font(r, size=9.3)


def build_docx():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_body_paragraph(doc, "Dear Dr. Zhang and Reviewers,")
    add_body_paragraph(
        doc,
        "Thank you for the careful evaluation and constructive comments. We have revised the manuscript comprehensively and provide point-by-point responses below. Reviewer comments are reproduced verbatim before each response. Where the manuscript was changed, the exact revised passage is reproduced in a highlighted box; additional experimental claims are accompanied by the corresponding figures or numerical evidence.",
    )

    p = doc.add_paragraph()
    r = p.add_run("Principal revisions")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    keep_with_next(p)
    bullets = [
        "redesigned Fig. 1 and clarified every region, link, and service symbol;",
        "renamed all deployment, provisioning, and association variables consistently;",
        "justified the two-stage formulation and quantified a small-scale exact joint-optimization gap;",
        "made the PSP capacity-repair mechanism explicit in Fig. 2, Algorithm 3, and the text;",
        "re-exported experimental figures with readable labels, ticks, and consistent panel dimensions;",
        "added CLS initialization sensitivity, complete PSP parameters, Q normalization, and Pareto metrics;",
        "included a DQN learning-based provisioning baseline; and",
        "added a real-region two-stage generalization experiment as response evidence."
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        r = p.add_run(item)
        set_run_font(r, size=10.2)

    current_section = None
    for item in ITEMS:
        if item["section"] != current_section:
            current_section = item["section"]
            doc.add_heading(f"Response to {current_section}", level=1)
        heading = "Associate Editor Comment" if item["number"] == "AE" else f"Comment {item['number']}: {item['title']}"
        doc.add_heading(heading, level=2)
        comment_label = "Associate Editor Comment" if item["number"] == "AE" else "Reviewer Comment"
        add_comment_box(doc, item["comment_en"], label=comment_label)
        add_response_label(doc)
        for paragraph in item["response_en"]:
            add_body_paragraph(doc, paragraph)
        excerpt = REVISED_EXCERPTS.get(item.get("evidence"))
        if excerpt:
            add_revised_text_box(doc, excerpt)
        if item.get("evidence"):
            add_evidence(doc, item["evidence"])

    doc.add_heading("Closing", level=1)
    add_body_paragraph(
        doc,
        "We appreciate the Editor's and Reviewers' comments, which have improved the clarity, reproducibility, and experimental support of the manuscript. We hope that the revised manuscript is now suitable for further consideration.",
    )
    add_body_paragraph(doc, "Sincerely,")
    add_body_paragraph(doc, "The Authors")

    doc.core_properties.title = "Response to the Editor and Reviewers - IoT-65990-2026"
    doc.core_properties.subject = "Point-by-point response and supporting evidence"
    doc.core_properties.author = "The Authors"
    doc.save(OUTPUT_DOCX)


def md_table(headers, rows):
    headers = [markdown_math(str(value)) for value in headers]
    rows = [[markdown_math(str(value)) for value in row] for row in rows]
    lines = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    lines.extend("| " + " | ".join(str(v) for v in row) + " |" for row in rows)
    return "\n".join(lines)


def excerpt_md(excerpt):
    lines = ["**论文修改后原文**", "", f"**位置：** {excerpt['location']}", ""]
    for block in excerpt["blocks"]:
        if block["type"] == "equation":
            lines.extend([f"$${FORMULA_LATEX[block['key']]}$$", ""])
        else:
            prefix = "*" if block["type"] == "caption" else ""
            lines.extend([f"> {prefix}{markdown_math(block['en'])}{prefix}", ""])
    lines.extend(["**修改后原文中文翻译**", ""])
    for block in excerpt["blocks"]:
        if block["type"] == "equation":
            lines.extend([f"$${FORMULA_LATEX[block['key']]}$$", block["zh"], ""])
        else:
            lines.extend([markdown_math(block["zh"]), ""])
    return "\n".join(lines)


def evidence_md(key):
    if key == "fig1":
        return "证据图：`response_evidence/fig1_revised.png`（修改后的 Fig. 1）。"
    if key == "notation":
        return md_table(["决策含义", "修改后变量", "定义"], [["服务器部署", "{{y_jk}}", "{{m_j}} 是否部署在 {{b_k}}"], ["服务配置", "{{z_jw}}", "{{m_j}} 是否配置 {{s_w}}"], ["用户关联", "{{a_ijw}}", "{{u_i}} 是否由 {{m_j}} 上的 {{s_w}} 服务"]])
    if key == "joint":
        return "证据图：`response_evidence/joint_exact_vs_mos2_seed42.png`。\n\n" + md_table(["种子", "HV 差距", "IGD", "Best Q 差距", "Joint/MOS\u00b2 运行时间"], joint_rows())
    if key == "fig2":
        return "证据图：`response_evidence/fig2_evolution_panel.png`（重绘的进化优化流程）。"
    if key == "readability":
        return "代表性重导出图：`response_evidence/stage1_scale.png`、`response_evidence/stage2_fixed_users_dqn.png`。"
    if key == "proofreading":
        return md_table(["原文", "修改后"], [["s2 , s3 and s5", "s2, s3, and s5"], ["U the set of", "U is the set of"], ["4rd / 5rd", "4th / 5th"], ["hierarchy network architecture", "hierarchical network architecture"]])
    if key == "cls":
        return "证据图：`response_evidence/cls_initialization_sensitivity.png`。"
    if key == "parameters":
        return md_table(["参数", "设置"], [["种群规模", "{{N50}}"], ["最大迭代代数", "{{G200}}"], ["SBX 交叉", "{{pc_eta}}"], ["多项式变异", "{{pm_eta}}"], ["混合评分权重", "{{alpha_beta}}"], ["服务容量", "{{Vj4}}"], ["锚点比例", "{{rho05}}"]])
    if key == "q":
        return ""
    if key == "varpi":
        return "证据图：`response_evidence/hybrid_initialization.png`。"
    if key == "generalization":
        return (
            "证据图：`response_evidence/real_region_topology.png`、`response_evidence/real_region_bestq.png`。\n\n"
            + md_table(["Method", "HV mean \u00b1 std", "IGD mean \u00b1 std", "Best Q mean \u00b1 std"], real_region_stage2_rows())
            + "\n\n"
            + md_table(["Method", "Best Q mean", "Std."], real_region_bestq_rows())
        )
    if key == "dqn":
        return "证据图：`response_evidence/stage2_fixed_servers_dqn.png`、`response_evidence/stage2_fixed_users_dqn.png`。\n\n" + md_table(["系列", "服务器/用户", "PSP Q", "DQN Q", "PSP 相对降低"], stage2_dqn_rows())
    if key == "pareto":
        return "证据图：`response_evidence/pareto_10_130.png`。\n\n" + md_table(["Method", "HV", "IGD", "Best Q"], [["NS-P", "0.8191", "0.0785", "0.3894"], ["GCP", "0.8596", "0.0492", "0.3550"], ["GDP", "0.8945", "0.0326", "0.3363"], ["PSP", "0.9470", "0.0016", "0.3282"], ["DQN", "N/A", "N/A", "0.6125"]])
    if key == "references":
        return "已引用建议文献 [1]、[3]、[4]、[5]；建议文献 [2]、[6] 因专门面向 FSO-SAGIN 任务卸载而未纳入。"
    return ""


def build_chinese_md():
    lines = [
        "# 回信逐条中英对照核对稿",
        "",
        "**原稿编号：** IoT-65990-2026  ",
        "**论文题目：** *MOS\u00b2: A Two-Stage Multi-Objective Framework for Server Deployment and Service Provisioning in Mobile Edge Computing*",
        "",
        "> 本文档与 `07_response_to_editor_and_reviewers.docx` 的条目顺序和回复内容一一对应，用于中文核对。Word 中的浅蓝色框直接列出修改后的论文原文；仅用于回信佐证的联合优化和真实区域实验不作为论文正文修改列出。",
        "",
    ]
    current_section = None
    section_zh = {"Associate Editor": "副编辑", "Reviewer 1": "审稿人 1", "Reviewer 2": "审稿人 2"}
    for item in ITEMS:
        if item["section"] != current_section:
            current_section = item["section"]
            lines.extend([f"## {section_zh[current_section]}", ""])
        title = "副编辑意见" if item["number"] == "AE" else f"### 意见 {item['number']}：{item['title']}"
        if item["number"] == "AE":
            lines.extend([f"### {title}", ""])
        else:
            lines.extend([title, ""])
        lines.extend(["**英文建议原文**", "", f"> {markdown_math(item['comment_en'])}", "", "**建议中文翻译**", "", markdown_math(item["comment_zh"]), "", "**英文回复**", ""])
        for p in item["response_en"]:
            lines.extend([markdown_math(p), ""])
        lines.extend(["**中文回复**", ""])
        for p in item["response_zh"]:
            lines.extend([markdown_math(p), ""])
        excerpt = REVISED_EXCERPTS.get(item.get("evidence"))
        if excerpt:
            lines.extend([excerpt_md(excerpt), ""])
        if item.get("evidence"):
            evidence = evidence_md(item["evidence"])
            if evidence:
                lines.extend(["**证据与数据**", "", evidence, ""])
        lines.extend(["---", ""])
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_chinese_md()
    print(OUTPUT_DOCX)
    print(OUTPUT_MD)
